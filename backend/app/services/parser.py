import pdfplumber
from docx import Document
from pathlib import Path
from typing import Optional


class DocumentParser:
    """Extract text from various document formats."""

    SUPPORTED_TYPES = {'.pdf', '.docx', '.doc', '.txt'}

    def parse(self, file_path: Path) -> str:
        """Main entry point - routes to appropriate parser."""
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return self._parse_pdf(file_path)
        elif suffix in {'.docx', '.doc'}:
            return self._parse_docx(file_path)
        elif suffix == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber for better accuracy."""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n\n".join(text_parts)

    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from Word documents."""
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables (common in syllabi)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    def _parse_txt(self, file_path: Path) -> str:
        """Read plain text file."""
        return file_path.read_text(encoding='utf-8')


# Singleton instance
parser = DocumentParser()
